# extractors.py
# ------------------------------------------------------------
# Tác dụng:
#   - Tự động đọc nội dung từ các định dạng tài liệu khác nhau (.docx, .txt, .pdf)
#   - Trả về text thuần để đưa vào embedding
# ------------------------------------------------------------

import os
from docx import Document
from pathlib import Path

try:
    import fitz  # PyMuPDF - đọc PDF
except ImportError:
    fitz = None


def read_txt(path):
    """Đọc file .txt"""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()


def read_docx(path):
    """Đọc file .docx"""
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def read_pdf(path):
    """Đọc file .pdf (nếu có PyMuPDF)"""
    if not fitz:
        raise ImportError("⚠️ Cần cài đặt PyMuPDF: pip install PyMuPDF")

    text = ""
    with fitz.open(path) as pdf:
        for page in pdf:
            text += page.get_text("text")
    return text.strip()


# extractors.py
from docx import Document
import os

def auto_extract(path: str) -> str:
    """Tự động trích xuất nội dung từ file docx hoặc txt, UTF-8."""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        try:
            doc = Document(path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            print(f"[!] Lỗi đọc file {path}: {e}")
            return ""
    elif ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        return ""
